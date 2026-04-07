import os
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from analysis import analyze_data


def _set_cell_bg(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def generate_word_report(file_path: str, output_path: str = "output_report.docx") -> str:
    try:
        df = pd.read_excel(file_path)
    except Exception as e:
        return f"Could not read file: {str(e)}"

    analysis = analyze_data(file_path)
    doc = Document()

    # ── Styles ───────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("GPT Excel — Data Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x1E, 0x27, 0x61)
    run.font.size = Pt(22)

    doc.add_paragraph()

    # ── File Info ─────────────────────────────────────────────────────────────
    info = doc.add_paragraph()
    info.add_run("File: ").bold = True
    info.add_run(os.path.basename(file_path))
    info.add_run("    |    ")
    info.add_run("Rows: ").bold = True
    info.add_run(str(analysis["total_rows"]))
    info.add_run("    |    ")
    info.add_run("Columns: ").bold = True
    info.add_run(str(analysis["total_columns"]))
    info.add_run("    |    ")
    info.add_run("Duplicates: ").bold = True
    info.add_run(str(analysis["duplicate_rows"]))

    doc.add_paragraph()

    # ── Column Overview ───────────────────────────────────────────────────────
    doc.add_heading("Column Overview", level=1)

    col_table = doc.add_table(rows=1, cols=4)
    col_table.style = "Table Grid"
    col_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Column Name", "Data Type", "Null Count", "Null %"]
    hdr_cells = col_table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _set_cell_bg(hdr_cells[i], "1E2761")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)

    for col in analysis["columns"]:
        row_cells = col_table.add_row().cells
        row_cells[0].text = str(col)
        row_cells[1].text = str(analysis["data_types"].get(col, ""))
        row_cells[2].text = str(analysis["null_values"].get(col, 0))
        row_cells[3].text = f"{analysis['null_percentage'].get(col, 0.0):.1f}%"
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── Statistics ────────────────────────────────────────────────────────────
    if isinstance(analysis.get("statistics"), dict):
        doc.add_heading("Statistical Summary", level=1)

        stat_keys = ["count", "mean", "std", "min", "25%", "50%", "75%", "max"]
        numeric_cols = list(analysis["statistics"].keys())

        stat_table = doc.add_table(rows=1, cols=len(numeric_cols) + 1)
        stat_table.style = "Table Grid"

        hdr = stat_table.rows[0].cells
        hdr[0].text = "Statistic"
        _set_cell_bg(hdr[0], "1E2761")
        hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[0].paragraphs[0].runs[0].font.bold = True

        for i, col in enumerate(numeric_cols):
            hdr[i + 1].text = col
            _set_cell_bg(hdr[i + 1], "4472C4")
            hdr[i + 1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            hdr[i + 1].paragraphs[0].runs[0].font.bold = True

        for sk in stat_keys:
            row = stat_table.add_row().cells
            row[0].text = sk
            row[0].paragraphs[0].runs[0].font.bold = True
            for i, col in enumerate(numeric_cols):
                val = analysis["statistics"][col].get(sk, "")
                try:
                    row[i + 1].text = f"{float(val):,.2f}"
                except Exception:
                    row[i + 1].text = str(val)
                row[i + 1].paragraphs[0].runs[0].font.size = Pt(10)

        doc.add_paragraph()

    # ── Data Preview ──────────────────────────────────────────────────────────
    doc.add_heading("Data Preview (First 10 Rows)", level=1)

    preview_df = df.head(10)
    max_cols = min(len(df.columns), 8)
    preview_cols = list(df.columns[:max_cols])

    preview_table = doc.add_table(rows=1, cols=len(preview_cols))
    preview_table.style = "Table Grid"

    ph = preview_table.rows[0].cells
    for i, col in enumerate(preview_cols):
        ph[i].text = str(col)
        _set_cell_bg(ph[i], "4472C4")
        ph[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        ph[i].paragraphs[0].runs[0].font.bold = True
        ph[i].paragraphs[0].runs[0].font.size = Pt(9)

    for _, data_row in preview_df.iterrows():
        row_cells = preview_table.add_row().cells
        for i, col in enumerate(preview_cols):
            row_cells[i].text = str(data_row[col])
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Generated by GPT Excel Python Engine")
    footer_run.font.color.rgb = RGBColor(0x88, 0x99, 0xBB)
    footer_run.font.size = Pt(9)
    footer_run.italic = True

    try:
        doc.save(output_path)
        return f"Word document generated: {output_path}"
    except Exception as e:
        return f"Could not save document: {str(e)}"
